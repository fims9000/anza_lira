#!/usr/bin/env python3
"""Create terminal claim-safe ANZA-LIRA manuscripts and reproducibility package."""

from __future__ import annotations

import hashlib
import json
from pathlib import Path
import zipfile

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parents[1]
DOWNLOADS = Path("/home/lebedeffson/Downloads")
RESULT = ROOT / "results/lira_h1/final"
ABSTRACT_SOURCE = DOWNLOADS / "MTUIP_2026_extended_abstract_UPDATED_20260819.docx"
FULL_SOURCE = DOWNLOADS / "ANZA_seismic_faults_full_research_UPDATED_20260819.docx"
ABSTRACT_TARGET = DOWNLOADS / "MTUIP_2026_extended_abstract_FINAL_ANZA_LIRA.docx"
FULL_TARGET = DOWNLOADS / "ANZA_LIRA_full_research_FINAL.docx"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def insert_before(anchor: Paragraph, text: str, style: str) -> Paragraph:
    element = OxmlElement("w:p")
    anchor._p.addprevious(element)
    paragraph = Paragraph(element, anchor._parent)
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def finalize_doc(source: Path, target: Path, paragraphs: list[tuple[str, str]]) -> None:
    document = Document(source)
    anchor = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip().upper() == "ЛИТЕРАТУРА")
    for style, text in paragraphs:
        insert_before(anchor, text, style)
    document.core_properties.title = "ANZA-LIRA: final frozen research report"
    document.core_properties.subject = "Final claim-safe closeout after H1"
    document.save(target)


def main() -> None:
    master = json.loads((RESULT / "ANZA_LIRA_H1_MASTER_RESULT.json").read_text())
    if master["status"] != "STOP_H1_RIBBON_BENCHMARK_FAIL":
        raise SystemExit("refusing to finalize documents for an unexpected H1 status")
    abstract_paragraphs = [
        ("AbsHeading", "ФИНАЛЬНАЯ ПРОВЕРКА ПЕРЕНОСА И ГРАНИЦЫ РЕЗУЛЬТАТА"),
        ("AbsBody", "Последняя correctness-проверка воспроизвела ошибку round-cap вмешательства: круглый торец действительно удалял ближайший видимый контекст. Точный flat-cap ribbon по проекции на полную полилинию сохранил anchors и прошёл геометрические unit tests. Однако на механическом CRACKS-аудите только 34 из 1753 предварительно допустимых локальных трасс удалось разъединить радиусами 3–15 пикселей; retention 0,0194 не достиг замороженного порога 0,50. Поэтому свежие секции 347–400, SBPP, P0, path и expert не открывались."),
        ("AbsBody", "Итоговая положительная часть имеет контролируемую область применимости: на независимой synthetic confirm классификатор пар достиг AUROC 0,9923, а learned completion восстановил 67,2 % скрытых разрывов при 0,78 % ложных мостов; oracle max-min восстановил все положительные разрывы без ложных связей. При этом learned recovery не достиг внутреннего порога 0,70. Устойчивого Anosov-specific преимущества и переноса на реальные естественные разрывы CRACKS не установлено."),
    ]
    full_paragraphs = [
        ("FullH1", "19. Финальная correctness-проверка H1"),
        ("FullBody", "Graph-Cut V2 использовал евклидову capsule вокруг скрытого участка. При радиусах 9–11 пикселей её круглые торцы заходили за границы скрытого интервала и удаляли обязательный видимый anchor. В H1 intervention был заменён на exact flat-cap ribbon: пиксель удаляется только тогда, когда его ближайшая проекция на полную ordered trace имеет arclength внутри замороженного hidden interval и поперечное расстояние не превосходит радиус. Unit tests воспроизвели старый дефект и подтвердили disconnection, anchor preservation, отсутствие longitudinal spillover, работу на кривой, reversal invariance и collateral rejection."),
        ("FullBody", "Исправление primitive не открыло новый performance experiment. На уже открытых секциях 263–344 механический audit имел 1753 pre-treatment eligible traces, но допустимый cut существовал лишь для 34 случаев: retention 0,019395 при замороженном gate 0,50. Для 1719 случаев связность на support P>=0,12 сохранялась при всех разрешённых радиусах до 15 пикселей. Поэтому статус H1 — STOP_H1_RIBBON_BENCHMARK_FAIL. Свежие секции 347–372 и final confirm 375–400 не читались; SBPP, P0, widest path и expert evaluation не запускались."),
        ("FullH1", "20. Итоговые утверждения и ограничения"),
        ("FullBody", "Наиболее сильный подтверждённый измерительный результат относится к independent controlled synthetic continuation: pair AUROC 0,9923, recovery 0,6719 и false-bridge 0,0078 при train-frozen threshold; oracle max-min дал recovery 1,0 без false links. Этот результат показывает потенциал разделения dense evidence, contextual relation verification и bounded reconstruction, но не является доказательством natural-gap recovery на CRACKS. Кроме того, learned recovery формально не прошёл заранее заданный gate 0,70, поэтому threshold не ретюнировался."),
        ("FullBody", "Серия causal controls не подтвердила устойчивого практически значимого incremental преимущества Anosov-specific локальных операторов. Корректный итог работы состоит не в постфактум выборе удачной ANZA-версии, а в зафиксированной цепочке фальсификаций и локализации уровней системы: локальная геометрия, candidate representation, relation calibration и path reconstruction требуют раздельной проверки. Новых rescue-веток после H1 не открывается."),
    ]
    finalize_doc(ABSTRACT_SOURCE, ABSTRACT_TARGET, abstract_paragraphs)
    finalize_doc(FULL_SOURCE, FULL_TARGET, full_paragraphs)
    # Store stable copies beside the final evidence.
    (RESULT / ABSTRACT_TARGET.name).write_bytes(ABSTRACT_TARGET.read_bytes())
    (RESULT / FULL_TARGET.name).write_bytes(FULL_TARGET.read_bytes())
    package = Path("/home/lebedeffson/Code/_wip_backups/anza_lira/ANZA_LIRA_H1_FINAL_RESEARCH_20260819.zip")
    package.parent.mkdir(parents=True, exist_ok=True)
    members = [
        ROOT / "results/lira_h1/freeze/protocol.json",
        ROOT / "results/lira_h1/freeze/split_manifest.json",
        ROOT / "results/lira_h1/freeze/H1_FRESH_SPLIT_AUTHORIZATION.json",
        ROOT / "results/lira_h1/ribbon_unit_tests/report.json",
        ROOT / "results/lira_h1/ribbon_unit_tests/H1_RIBBON_IMPLEMENTATION_REPORT.md",
        ROOT / "results/lira_h1/ribbon_unit_tests/figures/capsule_vs_ribbon.png",
        ROOT / "results/lira_h1/bug_audit/metrics.json",
        ROOT / "results/lira_h1/bug_audit/H1_BUG_AUDIT_REPORT.md",
        ROOT / "results/lira_h1/bug_audit/radius_distribution.csv",
        ROOT / "results/lira_h1/bug_audit/eligibility.csv",
        ROOT / "results/lira_h1/bug_audit/cases.csv",
        ROOT / "results/lira_h1/development_candidate/H1_FRESH_CANDIDATE_REPORT.md",
        *sorted(path for path in RESULT.iterdir() if path.name != "FINAL_RESEARCH_PACKAGE_MANIFEST.json"),
    ]
    unique = []
    for path in members:
        if path.is_file() and path not in unique:
            unique.append(path)
    with zipfile.ZipFile(package, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in unique:
            archive.write(path, path.relative_to(ROOT) if path.is_relative_to(ROOT) else Path("manuscripts") / path.name)
    manifest = {
        "status": master["status"],
        "abstract": {"path": str(ABSTRACT_TARGET), "sha256": sha256(ABSTRACT_TARGET)},
        "full_research": {"path": str(FULL_TARGET), "sha256": sha256(FULL_TARGET)},
        "package": {"path": str(package), "sha256": sha256(package)},
        "package_members": len(unique),
    }
    (RESULT / "FINAL_RESEARCH_PACKAGE_MANIFEST.json").write_text(json.dumps(manifest, indent=2, sort_keys=True) + "\n")
    print(json.dumps(manifest, indent=2, sort_keys=True))


if __name__ == "__main__":
    main()
